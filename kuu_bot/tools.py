import subprocess
import os
import glob as _glob
import re
import shutil
import zipfile
import urllib.request

DEFAULT_DIR = r"D:\妙妙工具\KuuBot"

TOOLS_DESC = f"""
当前环境: Windows 10/11, Python 3.10
默认工作目录: {DEFAULT_DIR}

⚠️ 核心规则：
1. read 前先 glob/bash 确认文件存在。
2. read 失败如实报告，绝对禁止编造内容。
3. 改文件前先 read 确认。
4. 修改后告诉主人改了什么。
"""

TOOL_DEFS = [
    {
        "type": "function",
        "function": {
            "name": "bash",
            "description": "执行 PowerShell 命令（不是 CMD！）。修改代码后必须执行测试命令检查是否正常运行。报错→分析→再改→再测。不要用 /s /b /a 等 CMD 标志",
            "parameters": {
                "type": "object",
                "properties": {"command": {"type": "string", "description": "要执行的命令"}},
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read",
            "description": f"读取文件内容。必须是绝对路径。默认目录: {DEFAULT_DIR}",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "文件绝对路径"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write",
            "description": "创建新文件或完全覆盖已有文件（⚠️ 全部覆盖！只改已有文件的某几行必须用 edit 而不是 write）",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "文件绝对路径"},
                    "content": {"type": "string", "description": "文件内容"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit",
            "description": "修改已有文件内容的推荐方式！只改指定文本不动其余部分。必须先 read 确认目标文本存在。改几行、改一个值、替换一段文字都必须用 edit。修改代码后必须重新测试，bash 返回错误→分析→再改→再测",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "文件绝对路径"},
                    "old": {"type": "string", "description": "要替换的文本"},
                    "new": {"type": "string", "description": "替换为的文本"},
                },
                "required": ["path", "old", "new"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "glob",
            "description": f"搜索匹配的文件名。path 默认 {DEFAULT_DIR}",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "文件匹配模式，如 *.py 或 **/*.json"},
                    "path": {"type": "string", "description": f"搜索目录，默认 {DEFAULT_DIR}"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "grep",
            "description": f"在文件内容中搜索正则匹配。path 默认 {DEFAULT_DIR}",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "正则表达式"},
                    "path": {"type": "string", "description": f"搜索目录，默认 {DEFAULT_DIR}"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "docx",
            "description": "创建 Word 文档(.docx)。必须是绝对路径",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "保存路径，如 C:/Users/31711/Desktop/报告.docx"},
                    "title": {"type": "string", "description": "文档标题"},
                    "content": {"type": "string", "description": "正文（支持 `# 标题` `**加粗**` `- 列表` 等简单 Markdown）"},
                },
                "required": ["path", "title", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete",
            "description": "删除文件或空目录。必须是绝对路径。不能删非空目录",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "要删除的文件/目录路径"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "mkdir",
            "description": f"创建目录。默认目录: {DEFAULT_DIR}",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "目录路径"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "copy",
            "description": "复制文件。source 和 dest 都必须是绝对路径",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "源文件路径"},
                    "dest": {"type": "string", "description": "目标路径"},
                },
                "required": ["source", "dest"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "clipboard",
            "description": "读取 Windows 剪贴板文本内容",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "notify",
            "description": "弹出 Windows 系统通知。用于提醒主人",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "通知标题"},
                    "msg": {"type": "string", "description": "通知内容"},
                },
                "required": ["title", "msg"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "sysinfo",
            "description": "获取系统信息（磁盘空间、内存使用等）",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "screenshot",
            "description": "全屏截图并保存为 PNG文件。必须是绝对路径",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "保存路径，如 C:/Users/31711/Desktop/screenshot.png"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "zip",
            "description": "压缩文件或目录为 zip。source 和 dest 都必须是绝对路径",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "要压缩的文件/目录路径"},
                    "dest": {"type": "string", "description": "zip 文件保存路径"},
                },
                "required": ["source", "dest"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "unzip",
            "description": "解压 zip 文件到指定目录。source 和 dest 都必须是绝对路径",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "zip 文件路径"},
                    "dest": {"type": "string", "description": "解压目标目录"},
                },
                "required": ["source", "dest"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "download",
            "description": "从 URL 下载文件。url 是下载地址，path 是保存的绝对路径",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "下载链接"},
                    "path": {"type": "string", "description": "保存的绝对路径"},
                },
                "required": ["url", "path"],
            },
        },
    },
]


def execute(tool_name: str, args: dict) -> str:
    if tool_name == "bash":
        cmd = args.get("command", "")
        try:
            r = subprocess.run(
                ["powershell", "-Command", cmd],
                capture_output=True, text=True, timeout=60,
                cwd=DEFAULT_DIR,
            )
            out = r.stdout.strip() or r.stderr.strip()
            return out[:4000] if out else "(no output)"
        except subprocess.TimeoutExpired:
            return "(timeout)"
        except Exception as e:
            return str(e)

    elif tool_name == "read":
        path = args.get("path", "")
        if not os.path.isabs(path):
            return f"read 失败：路径必须是绝对路径，不能是相对路径。当前默认目录: {DEFAULT_DIR}"
        try:
            if path.lower().endswith(".docx"):
                from docx import Document
                doc = Document(path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n".join(paragraphs)[:8000]
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()[:8000]
        except FileNotFoundError:
            return f"read 失败：文件不存在 [{path}]，请先用 glob 搜索确认路径"
        except PermissionError:
            return f"read 失败：权限不足 [{path}]"
        except Exception as e:
            return f"read 失败 [{path}]: {e}"

    elif tool_name == "write":
        path = args.get("path", "")
        content = args.get("content", "")
        if not os.path.isabs(path):
            return f"write 失败：路径必须是绝对路径。默认目录: {DEFAULT_DIR}"
        try:
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            return f"写入成功: {path} ({len(content)} bytes)"
        except PermissionError:
            return f"write 失败：权限不足 [{path}]（D盘根目录需要管理员权限）"
        except Exception as e:
            return f"write 失败 [{path}]: {e}"

    elif tool_name == "edit":
        path = args.get("path", "")
        old = args.get("old", "")
        new = args.get("new", "")
        if not os.path.isabs(path):
            return f"edit 失败：路径必须是绝对路径。默认目录: {DEFAULT_DIR}"
        if not old or not new:
            return "edit: old and new strings required"
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            if old not in text:
                return f"edit 失败 [{path}]: 未找到指定文本"
            text = text.replace(old, new, 1)
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            return f"修改成功: {path}"
        except FileNotFoundError:
            return f"edit 失败：文件不存在 [{path}]"
        except Exception as e:
            return f"edit 失败 [{path}]: {e}"

    elif tool_name == "glob":
        pattern = args.get("pattern", "**/*")
        path = args.get("path", DEFAULT_DIR)
        try:
            matches = _glob.glob(pattern, root_dir=path, recursive=True)[:50]
            if not matches:
                return "(no matches)"
            return "\n".join(os.path.join(path, m) for m in matches)
        except Exception as e:
            return str(e)

    elif tool_name == "grep":
        pattern = args.get("pattern", "")
        path = args.get("path", DEFAULT_DIR)
        try:
            results = []
            for root, dirs, files in os.walk(path):
                dirs[:] = [d for d in dirs if not d.startswith(".") and d not in ("node_modules", "__pycache__")]
                for fn in files[:100]:
                    if fn.startswith("."):
                        continue
                    fp = os.path.join(root, fn)
                    try:
                        with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                            for i, line in enumerate(f, 1):
                                if re.search(pattern, line):
                                    results.append(f"{fp}:{i}: {line.rstrip()[:200]}")
                                    if len(results) >= 30:
                                        break
                    except:
                        pass
                    if len(results) >= 30:
                        break
                if len(results) >= 30:
                    break
            return "\n".join(results) if results else "(no matches)"
        except Exception as e:
            return str(e)

    elif tool_name == "docx":
        path = args.get("path", "")
        title = args.get("title", "未命名")
        content = args.get("content", "")
        if not os.path.isabs(path):
            return f"docx 失败：路径必须是绝对路径。默认目录: {DEFAULT_DIR}"
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles["Normal"]
            style.font.size = Pt(11)
            style.font.name = "微软雅黑"

            # Title
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Parse simple markdown
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("1. "):
                    doc.add_paragraph(line[3:], style="List Number")
                else:
                    p = doc.add_paragraph()
                    # Handle **bold** inline
                    import re as _re
                    parts = _re.split(r"(\*\*.*?\*\*)", line)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            doc.save(path)
            return f"Word 文档创建成功: {path} (标题: {title})"
        except PermissionError:
            return f"docx 失败：权限不足 [{path}]"
        except Exception as e:
            return f"docx 失败 [{path}]: {e}"

    elif tool_name == "delete":
        path = args.get("path", "")
        if not os.path.isabs(path):
            return f"delete 失败：必须是绝对路径"
        try:
            if os.path.isdir(path):
                os.rmdir(path)
            else:
                os.remove(path)
            return f"已删除: {path}"
        except OSError as e:
            return f"delete 失败 [{path}]: {e}（试试 bash 强制删除？）"
        except Exception as e:
            return f"delete 失败 [{path}]: {e}"

    elif tool_name == "mkdir":
        path = args.get("path", "")
        if not os.path.isabs(path):
            path = os.path.join(DEFAULT_DIR, path)
        try:
            os.makedirs(path, exist_ok=True)
            return f"目录已创建: {path}"
        except Exception as e:
            return f"mkdir 失败 [{path}]: {e}"

    elif tool_name == "copy":
        source = args.get("source", "")
        dest = args.get("dest", "")
        if not os.path.isabs(source) or not os.path.isabs(dest):
            return "copy 失败：source 和 dest 必须是绝对路径"
        try:
            os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
            shutil.copy2(source, dest)
            return f"已复制: {source} -> {dest}"
        except Exception as e:
            return f"copy 失败: {e}"

    elif tool_name == "clipboard":
        try:
            import tkinter as tk
            root = tk.Tk()
            root.withdraw()
            text = root.clipboard_get()
            root.destroy()
            return text[:4000] if text.strip() else "(剪贴板为空)"
        except Exception as e:
            return f"clipboard 失败: {e}"

    elif tool_name == "notify":
        title = args.get("title", "KuuBot")
        msg = args.get("msg", "")
        try:
            import ctypes
            ctypes.windll.user32.MessageBoxW(0, msg, title, 0x40)
            return "通知已弹出"
        except Exception as e:
            return f"notify 失败: {e}"

    elif tool_name == "sysinfo":
        try:
            import psutil
            disk = psutil.disk_usage("C:")
            mem = psutil.virtual_memory()
            return (
                f"C盘: {disk.free//(1024**3)}GB 可用 / {disk.total//(1024**3)}GB 总计 ({disk.percent}%)\n"
                f"内存: {mem.available//(1024**2)}MB 可用 / {mem.total//(1024**2)}MB 总计 ({mem.percent}%)"
            )
        except ImportError:
            try:
                r = subprocess.run(["powershell", "-Command",
                    "Get-PSDrive C | Select @{N='FreeGB';E={$_.Free/1GB}},@{N='UsedGB';E={$_.Used/1GB}}"],
                    capture_output=True, text=True, timeout=10)
                return r.stdout.strip()[:2000] or "无法获取系统信息"
            except:
                return "无法获取系统信息"

    elif tool_name == "screenshot":
        path = args.get("path", "")
        if not os.path.isabs(path):
            return f"screenshot 失败：必须是绝对路径"
        try:
            import ctypes
            from ctypes import wintypes
            u32 = ctypes.windll.user32
            gdi = ctypes.windll.gdi32
            w = u32.GetSystemMetrics(0)
            h = u32.GetSystemMetrics(1)
            hdc_s = u32.GetDC(None)
            hdc_m = gdi.CreateCompatibleDC(hdc_s)
            hbmp = gdi.CreateCompatibleBitmap(hdc_s, w, h)
            gdi.SelectObject(hdc_m, hbmp)
            gdi.BitBlt(hdc_m, 0, 0, w, h, hdc_s, 0, 0, 0x00CC0020)
            from PIL import Image
            bmp_data = ctypes.create_string_buffer(w * h * 4)
            bmi = ctypes.create_string_buffer(ctypes.sizeof(wintypes.BITMAPINFOHEADER))
            bmi_hdr = ctypes.cast(bmi, ctypes.POINTER(wintypes.BITMAPINFOHEADER))
            bmi_hdr[0].biSize = ctypes.sizeof(wintypes.BITMAPINFOHEADER)
            bmi_hdr[0].biWidth = w
            bmi_hdr[0].biHeight = -h
            bmi_hdr[0].biPlanes = 1
            bmi_hdr[0].biBitCount = 32
            gdi.GetDIBits(hdc_m, hbmp, 0, h, bmp_data, bmi, 0)
            img = Image.frombuffer("RGB", (w, h), bmp_data, "raw", "BGRX", 0, 1)
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            img.save(path, "PNG")
            gdi.DeleteObject(hbmp)
            gdi.DeleteDC(hdc_m)
            u32.ReleaseDC(None, hdc_s)
            return f"截图已保存: {path} ({w}x{h})"
        except ImportError:
            return "screenshot 需要 PIL: pip install Pillow"
        except Exception as e:
            return f"screenshot 失败: {e}"

    elif tool_name == "zip":
        source = args.get("source", "")
        dest = args.get("dest", "")
        if not os.path.isabs(source) or not os.path.isabs(dest):
            return "zip 失败：source 和 dest 必须是绝对路径"
        try:
            os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
            if os.path.isdir(source):
                with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(source):
                        for fn in files:
                            fp = os.path.join(root, fn)
                            zf.write(fp, os.path.relpath(fp, source))
            else:
                with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.write(source, os.path.basename(source))
            size = os.path.getsize(dest)
            return f"压缩完成: {dest} ({size} bytes)"
        except Exception as e:
            return f"zip 失败: {e}"

    elif tool_name == "unzip":
        source = args.get("source", "")
        dest = args.get("dest", "")
        if not os.path.isabs(source) or not os.path.isabs(dest):
            return "unzip 失败：source 和 dest 必须是绝对路径"
        try:
            os.makedirs(dest, exist_ok=True)
            with zipfile.ZipFile(source, "r") as zf:
                zf.extractall(dest)
            return f"解压完成: {dest} ({len(zf.namelist())} 个文件)"
        except Exception as e:
            return f"unzip 失败: {e}"

    elif tool_name == "download":
        url = args.get("url", "")
        path = args.get("path", "")
        if not os.path.isabs(path):
            return f"download 失败：路径必须是绝对路径"
        try:
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            urllib.request.urlretrieve(url, path)
            size = os.path.getsize(path)
            return f"下载完成: {path} ({size} bytes)"
        except Exception as e:
            return f"download 失败: {e}"

    else:
        return f"Unknown tool: {tool_name}"
